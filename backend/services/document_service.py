import pytesseract
import pdfplumber
from PIL import Image
from docx import Document
import os

class DocumentService:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.jpg', '.jpeg', '.png']
    
    def extract_text(self, file_path: str) -> str:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_docx_text(file_path)
            elif file_extension in ['.jpg', '.jpeg', '.png']:
                return self._extract_image_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        
        except Exception as e:
            raise Exception(f"Failed to extract text from document: {str(e)}")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    
                    if not page_text and page.images:
                        for image in page.images:
                            try:
                                bbox = (image['x0'], image['top'], image['x1'], image['bottom'])
                                cropped_page = page.crop(bbox)
                                page_image = cropped_page.to_image()
                                pil_image = page_image.original
                                ocr_text = pytesseract.image_to_string(pil_image)
                                text += ocr_text + "\n"
                            except:
                                continue
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
        
        return text.strip()
    
    def _extract_docx_text(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def _extract_image_text(self, file_path: str) -> str:
        try:
            image = Image.open(file_path)
            
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            text = pytesseract.image_to_string(image, config='--psm 6')
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Error processing image: {str(e)}")
    
    def validate_document_structure(self, text: str) -> dict:
        validation_results = {
            'has_signature_section': False,
            'has_date_fields': False,
            'has_contact_info': False,
            'has_project_description': False,
            'has_dimensions': False,
            'word_count': len(text.split()),
            'issues': []
        }
        
        text_lower = text.lower()
        
        signature_keywords = ['signature', 'signed by', 'applicant signature', 'owner signature']
        if any(keyword in text_lower for keyword in signature_keywords):
            validation_results['has_signature_section'] = True
        else:
            validation_results['issues'].append("No signature section found")
        
        date_patterns = ['date:', '/20', '-20', 'dated']
        if any(pattern in text_lower for pattern in date_patterns):
            validation_results['has_date_fields'] = True
        else:
            validation_results['issues'].append("No date information found")
        
        contact_keywords = ['phone', 'email', 'address', 'contact']
        if any(keyword in text_lower for keyword in contact_keywords):
            validation_results['has_contact_info'] = True
        else:
            validation_results['issues'].append("Missing contact information")
        
        description_keywords = ['project description', 'scope of work', 'construction', 'building']
        if any(keyword in text_lower for keyword in description_keywords):
            validation_results['has_project_description'] = True
        else:
            validation_results['issues'].append("Missing project description")
        
        dimension_patterns = ["ft", "feet", "inches", "'", '"', "x", "square feet", "sq ft"]
        if any(pattern in text_lower for pattern in dimension_patterns):
            validation_results['has_dimensions'] = True
        else:
            validation_results['issues'].append("Missing dimensional information")
        
        return validation_results