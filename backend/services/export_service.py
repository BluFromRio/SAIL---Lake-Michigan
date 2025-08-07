from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
import tempfile
import os
from typing import Dict, Any
from datetime import datetime

class ExportService:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            textColor=colors.HexColor('#1f2937')
        )
        self.heading_style = ParagraphStyle(
            'CustomHeading',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.HexColor('#374151')
        )
    
    def create_document(self, export_data: Dict[str, Any], export_type: str) -> str:
        if export_type == 'pdf':
            return self._create_pdf(export_data)
        elif export_type == 'docx':
            return self._create_docx(export_data)
        elif export_type == 'checklist':
            return self._create_checklist_pdf(export_data)
        else:
            raise ValueError(f"Unsupported export type: {export_type}")
    
    def _create_pdf(self, export_data: Dict[str, Any]) -> str:
        temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        doc = SimpleDocTemplate(temp_file.name, pagesize=letter, topMargin=0.75*inch)
        story = []
        
        story.append(Paragraph("🏗️ PermitCheck AI Report", self.title_style))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", self.styles['Normal']))
        story.append(Spacer(1, 20))
        
        project_data = export_data.get('project_data', {})
        if project_data:
            story.append(Paragraph("Project Information", self.heading_style))
            project_info = [
                ['Description:', project_data.get('description', 'N/A')],
                ['Address:', project_data.get('address', 'N/A')],
                ['Structure Type:', project_data.get('structure_type', 'N/A')],
                ['Property Type:', project_data.get('property_type', 'N/A')]
            ]
            
            dimensions = project_data.get('dimensions', {})
            if dimensions:
                dim_text = f"{dimensions.get('length', 'N/A')}' x {dimensions.get('width', 'N/A')}' x {dimensions.get('height', 'N/A')}'"
                project_info.append(['Dimensions:', dim_text])
            
            table = Table(project_info, colWidths=[1.5*inch, 4*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (0,-1), colors.HexColor('#f3f4f6')),
                ('TEXTCOLOR', (0,0), (-1,-1), colors.black),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,-1), 10),
                ('GRID', (0,0), (-1,-1), 1, colors.HexColor('#e5e7eb'))
            ]))
            story.append(table)
            story.append(Spacer(1, 20))
        
        feasibility_results = export_data.get('feasibility_results', {})
        if feasibility_results:
            story.append(Paragraph("Feasibility Assessment", self.heading_style))
            verdict = feasibility_results.get('verdict', 'Unknown')
            confidence = feasibility_results.get('confidence_score', 0)
            story.append(Paragraph(f"<b>Verdict:</b> {verdict} (Confidence: {confidence}%)", self.styles['Normal']))
            
            compliance_summary = feasibility_results.get('compliance_summary', '')
            if compliance_summary:
                story.append(Spacer(1, 10))
                story.append(Paragraph("<b>Compliance Summary:</b>", self.styles['Normal']))
                story.append(Paragraph(compliance_summary, self.styles['Normal']))
            
            issues = feasibility_results.get('issues', [])
            if issues:
                story.append(Spacer(1, 10))
                story.append(Paragraph("<b>Issues Identified:</b>", self.styles['Normal']))
                for issue in issues:
                    story.append(Paragraph(f"• {issue}", self.styles['Normal']))
            
            recommendations = feasibility_results.get('recommendations', [])
            if recommendations:
                story.append(Spacer(1, 10))
                story.append(Paragraph("<b>Recommendations:</b>", self.styles['Normal']))
                for rec in recommendations:
                    story.append(Paragraph(f"• {rec}", self.styles['Normal']))
            
            story.append(Spacer(1, 20))
        
        narrative_results = export_data.get('narrative_results', {})
        if narrative_results:
            story.append(Paragraph("Construction Narrative", self.heading_style))
            narrative = narrative_results.get('narrative', '')
            story.append(Paragraph(narrative, self.styles['Normal']))
            story.append(Spacer(1, 20))
        
        review_results = export_data.get('review_results', {})
        if review_results:
            story.append(PageBreak())
            story.append(Paragraph("Document Review Results", self.heading_style))
            
            rejection_risk = review_results.get('rejection_risk', 'Unknown')
            story.append(Paragraph(f"<b>Rejection Risk:</b> {rejection_risk}", self.styles['Normal']))
            
            overall_assessment = review_results.get('overall_assessment', '')
            if overall_assessment:
                story.append(Spacer(1, 10))
                story.append(Paragraph("<b>Overall Assessment:</b>", self.styles['Normal']))
                story.append(Paragraph(overall_assessment, self.styles['Normal']))
            
            issues = review_results.get('issues', [])
            if issues:
                story.append(Spacer(1, 15))
                story.append(Paragraph("<b>Issues Found:</b>", self.styles['Normal']))
                for issue in issues:
                    issue_text = issue if isinstance(issue, str) else issue.get('description', 'Unknown issue')
                    story.append(Paragraph(f"• {issue_text}", self.styles['Normal']))
            
            fixes = review_results.get('fixes', [])
            if fixes:
                story.append(Spacer(1, 15))
                story.append(Paragraph("<b>Recommended Fixes:</b>", self.styles['Normal']))
                for fix in fixes:
                    fix_text = fix if isinstance(fix, str) else fix.get('description', 'Unknown fix')
                    story.append(Paragraph(f"• {fix_text}", self.styles['Normal']))
        
        story.append(Spacer(1, 30))
        story.append(Paragraph("Generated by PermitCheck AI - AI-Powered Feasibility and Permit Review Assistant", self.styles['Normal']))
        
        doc.build(story)
        return temp_file.name
    
    def _create_docx(self, export_data: Dict[str, Any]) -> str:
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc = Document()
        
        title = doc.add_heading('PermitCheck AI Report', 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph()
        
        project_data = export_data.get('project_data', {})
        if project_data:
            doc.add_heading('Project Information', level=1)
            doc.add_paragraph(f"Description: {project_data.get('description', 'N/A')}")
            doc.add_paragraph(f"Address: {project_data.get('address', 'N/A')}")
            doc.add_paragraph(f"Structure Type: {project_data.get('structure_type', 'N/A')}")
            doc.add_paragraph(f"Property Type: {project_data.get('property_type', 'N/A')}")
            
            dimensions = project_data.get('dimensions', {})
            if dimensions:
                dim_text = f"{dimensions.get('length', 'N/A')}' x {dimensions.get('width', 'N/A')}' x {dimensions.get('height', 'N/A')}'"
                doc.add_paragraph(f"Dimensions: {dim_text}")
            doc.add_paragraph()
        
        feasibility_results = export_data.get('feasibility_results', {})
        if feasibility_results:
            doc.add_heading('Feasibility Assessment', level=1)
            verdict = feasibility_results.get('verdict', 'Unknown')
            confidence = feasibility_results.get('confidence_score', 0)
            doc.add_paragraph(f"Verdict: {verdict} (Confidence: {confidence}%)")
            
            compliance_summary = feasibility_results.get('compliance_summary', '')
            if compliance_summary:
                doc.add_paragraph()
                doc.add_paragraph("Compliance Summary:")
                doc.add_paragraph(compliance_summary)
            
            issues = feasibility_results.get('issues', [])
            if issues:
                doc.add_paragraph()
                doc.add_paragraph("Issues Identified:")
                for issue in issues:
                    doc.add_paragraph(f"• {issue}")
            
            recommendations = feasibility_results.get('recommendations', [])
            if recommendations:
                doc.add_paragraph()
                doc.add_paragraph("Recommendations:")
                for rec in recommendations:
                    doc.add_paragraph(f"• {rec}")
            doc.add_paragraph()
        
        narrative_results = export_data.get('narrative_results', {})
        if narrative_results:
            doc.add_heading('Construction Narrative', level=1)
            narrative = narrative_results.get('narrative', '')
            doc.add_paragraph(narrative)
            doc.add_paragraph()
        
        review_results = export_data.get('review_results', {})
        if review_results:
            doc.add_page_break()
            doc.add_heading('Document Review Results', level=1)
            
            rejection_risk = review_results.get('rejection_risk', 'Unknown')
            doc.add_paragraph(f"Rejection Risk: {rejection_risk}")
            
            overall_assessment = review_results.get('overall_assessment', '')
            if overall_assessment:
                doc.add_paragraph()
                doc.add_paragraph("Overall Assessment:")
                doc.add_paragraph(overall_assessment)
            
            issues = review_results.get('issues', [])
            if issues:
                doc.add_paragraph()
                doc.add_paragraph("Issues Found:")
                for issue in issues:
                    issue_text = issue if isinstance(issue, str) else issue.get('description', 'Unknown issue')
                    doc.add_paragraph(f"• {issue_text}")
            
            fixes = review_results.get('fixes', [])
            if fixes:
                doc.add_paragraph()
                doc.add_paragraph("Recommended Fixes:")
                for fix in fixes:
                    fix_text = fix if isinstance(fix, str) else fix.get('description', 'Unknown fix')
                    doc.add_paragraph(f"• {fix_text}")
        
        doc.add_paragraph()
        doc.add_paragraph("Generated by PermitCheck AI - AI-Powered Feasibility and Permit Review Assistant")
        
        doc.save(temp_file.name)
        return temp_file.name
    
    def _create_checklist_pdf(self, export_data: Dict[str, Any]) -> str:
        temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        doc = SimpleDocTemplate(temp_file.name, pagesize=letter, topMargin=0.75*inch)
        story = []
        
        story.append(Paragraph("📋 Permit Application Fix Checklist", self.title_style))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", self.styles['Normal']))
        story.append(Spacer(1, 20))
        
        review_results = export_data.get('review_results', {})
        if review_results:
            rejection_risk = review_results.get('rejection_risk', 'Unknown')
            story.append(Paragraph(f"Current Rejection Risk: <b>{rejection_risk}</b>", self.styles['Normal']))
            story.append(Spacer(1, 20))
            
            fixes = review_results.get('fixes', [])
            if fixes:
                story.append(Paragraph("Action Items:", self.heading_style))
                
                checklist_data = [['☐', 'Priority', 'Action Item']]
                for fix in fixes:
                    if isinstance(fix, dict):
                        priority = fix.get('priority', 'Medium')
                        description = fix.get('description', 'Unknown fix')
                    else:
                        priority = 'Medium'
                        description = fix
                    
                    checklist_data.append(['☐', priority, description])
                
                table = Table(checklist_data, colWidths=[0.3*inch, 0.8*inch, 4.5*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1f2937')),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0,0), (-1,0), 12),
                    ('BOTTOMPADDING', (0,0), (-1,0), 12),
                    ('BACKGROUND', (0,1), (-1,-1), colors.beige),
                    ('GRID', (0,0), (-1,-1), 1, colors.black),
                    ('FONTSIZE', (0,1), (-1,-1), 10),
                    ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')])
                ]))
                story.append(table)
            
            missing_docs = review_results.get('missing_documents', [])
            if missing_docs:
                story.append(Spacer(1, 30))
                story.append(Paragraph("Missing Documents:", self.heading_style))
                for doc in missing_docs:
                    story.append(Paragraph(f"☐ {doc}", self.styles['Normal']))
        
        story.append(Spacer(1, 40))
        story.append(Paragraph("Instructions:", self.heading_style))
        story.append(Paragraph("1. Check off each item as you complete it", self.styles['Normal']))
        story.append(Paragraph("2. Focus on High priority items first", self.styles['Normal']))
        story.append(Paragraph("3. Re-upload your revised document for another review", self.styles['Normal']))
        story.append(Paragraph("4. Repeat until rejection risk is Low", self.styles['Normal']))
        
        doc.build(story)
        return temp_file.name